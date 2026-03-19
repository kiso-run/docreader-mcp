#!/usr/bin/env python3
"""One-time script to generate static test fixture files.

Run: uv run python tests/create_fixtures.py
Requires: fpdf2 (temporary, `uv pip install fpdf2`)

Generated files are committed to the repo — this script only needs
to be re-run if fixtures need updating.
"""
from __future__ import annotations

import csv
from pathlib import Path

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)


def create_pdf():
    """Create a 3-page PDF with real embedded text."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)

    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Sample Document - Page 1")
    pdf.ln(10)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(w=0, text=(
        "This is the first page of the sample PDF document. "
        "It contains introductory text that can be extracted by pypdf. "
        "The document has three pages total."
    ))

    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Data Section - Page 2")
    pdf.ln(10)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(w=0, text=(
        "Page two contains the main data section. "
        "Server IP: 192.168.1.100. Database port: 5432. "
        "The deployment uses PostgreSQL 15 with pgvector extension."
    ))

    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Conclusion - Page 3")
    pdf.ln(10)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(w=0, text=(
        "This is the final page with conclusions. "
        "All systems are operational and ready for production deployment."
    ))

    pdf.output(str(FIXTURES / "sample.pdf"))
    print(f"  Created sample.pdf ({(FIXTURES / 'sample.pdf').stat().st_size} bytes)")


def create_docx():
    """Create a DOCX with heading + 2 paragraphs."""
    from docx import Document

    doc = Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph(
        "This project uses Flask as the web framework with SQLAlchemy for database access. "
        "The API endpoints follow REST conventions."
    )
    doc.add_paragraph(
        "Deployment is handled via Docker containers orchestrated by docker-compose. "
        "The CI pipeline runs on GitHub Actions."
    )
    doc.save(str(FIXTURES / "sample.docx"))
    print(f"  Created sample.docx ({(FIXTURES / 'sample.docx').stat().st_size} bytes)")


def create_xlsx():
    """Create an XLSX with 2 sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sales"
    ws1.append(["Date", "Product", "Amount", "Region"])
    ws1.append(["2024-01-15", "Widget A", 150, "Europe"])
    ws1.append(["2024-01-16", "Widget B", 230, "Asia"])
    ws1.append(["2024-01-17", "Widget A", 180, "Europe"])
    ws1.append(["2024-02-01", "Widget C", 95, "Americas"])

    ws2 = wb.create_sheet("Summary")
    ws2.append(["Region", "Total"])
    ws2.append(["Europe", 330])
    ws2.append(["Asia", 230])

    wb.save(str(FIXTURES / "sample.xlsx"))
    print(f"  Created sample.xlsx ({(FIXTURES / 'sample.xlsx').stat().st_size} bytes)")


def create_csv():
    """Create a CSV with header + 5 rows."""
    path = FIXTURES / "sample.csv"
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "name", "email", "role", "active"])
        writer.writerow([1, "Alice", "alice@example.com", "admin", "true"])
        writer.writerow([2, "Bob", "bob@example.com", "user", "true"])
        writer.writerow([3, "Carol", "carol@example.com", "user", "false"])
        writer.writerow([4, "Dave", "dave@example.com", "admin", "true"])
        writer.writerow([5, "Eve", "eve@example.com", "user", "true"])
    print(f"  Created sample.csv ({path.stat().st_size} bytes)")


def create_txt():
    """Create a plain text file with 10 lines."""
    path = FIXTURES / "sample.txt"
    lines = [
        "Application Configuration Notes",
        "================================",
        "",
        "Server: kiso-prod-01.example.com",
        "Port: 8334",
        "Database: PostgreSQL 15 on localhost:5432",
        "Cache: Redis on localhost:6379",
        "",
        "Last updated: 2024-03-15",
        "Maintainer: ops-team@example.com",
    ]
    path.write_text("\n".join(lines) + "\n")
    print(f"  Created sample.txt ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    print("Generating test fixtures...")
    create_pdf()
    create_docx()
    create_xlsx()
    create_csv()
    create_txt()
    print("Done.")
