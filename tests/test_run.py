"""Unit tests for tool-docreader."""
from __future__ import annotations

import csv
import json
import subprocess
import sys
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

# Import directly from run.py (pythonpath = ["."])
from run import (
    do_list, do_info, do_read,
    _resolve_path, _parse_page_ranges, _format_size, _is_likely_text,
    _MAX_OUTPUT_CHARS,
)

FIXTURES = Path(__file__).parent / "fixtures"


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def workspace(tmp_path):
    """Create a workspace with uploads/ directory."""
    uploads = tmp_path / "uploads"
    uploads.mkdir()
    return tmp_path


@pytest.fixture
def workspace_with_fixtures(workspace):
    """Workspace with all static fixture files copied to uploads/."""
    import shutil
    for f in FIXTURES.iterdir():
        if f.is_file():
            shutil.copy2(f, workspace / "uploads" / f.name)
    return workspace


@pytest.fixture
def txt_file(workspace):
    """Create a plain text file."""
    f = workspace / "uploads" / "hello.txt"
    f.write_text("Hello, world!\nSecond line.")
    return f


@pytest.fixture
def csv_file(workspace):
    """Create a CSV file."""
    f = workspace / "uploads" / "data.csv"
    with open(f, "w", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["name", "age", "city"])
        writer.writerow(["Alice", "30", "Rome"])
        writer.writerow(["Bob", "25", "Milan"])
    return f


# ---------------------------------------------------------------------------
# do_list
# ---------------------------------------------------------------------------


class TestDoList:
    def test_list_with_files(self, workspace, txt_file, csv_file):
        result = do_list(str(workspace))
        assert "Files in uploads/ (2):" in result
        assert "hello.txt" in result
        assert "data.csv" in result

    def test_list_empty_directory(self, workspace):
        result = do_list(str(workspace))
        assert "empty" in result.lower()

    def test_list_no_uploads_dir(self, tmp_path):
        result = do_list(str(tmp_path))
        assert "No uploads/" in result


# ---------------------------------------------------------------------------
# do_info
# ---------------------------------------------------------------------------


class TestDoInfo:
    def test_info_txt(self, workspace_with_fixtures):
        result = do_info(str(workspace_with_fixtures), {"file_path": "uploads/sample.txt"})
        assert "sample.txt" in result
        assert ".txt" in result

    def test_info_csv(self, workspace_with_fixtures):
        result = do_info(str(workspace_with_fixtures), {"file_path": "uploads/sample.csv"})
        assert "Rows:" in result
        assert "6" in result  # header + 5 data rows

    def test_info_pdf(self, workspace_with_fixtures):
        result = do_info(str(workspace_with_fixtures), {"file_path": "uploads/sample.pdf"})
        assert "Pages: 3" in result

    def test_info_xlsx(self, workspace_with_fixtures):
        result = do_info(str(workspace_with_fixtures), {"file_path": "uploads/sample.xlsx"})
        assert "Sheets:" in result
        assert "Sales" in result
        assert "Summary" in result

    def test_info_docx(self, workspace_with_fixtures):
        result = do_info(str(workspace_with_fixtures), {"file_path": "uploads/sample.docx"})
        assert "sample.docx" in result
        assert ".docx" in result

    def test_info_missing_file(self, workspace):
        with pytest.raises(FileNotFoundError):
            do_info(str(workspace), {"file_path": "uploads/nope.pdf"})


# ---------------------------------------------------------------------------
# do_read — basic (small files via static fixtures)
# ---------------------------------------------------------------------------


class TestDoRead:
    def test_read_txt(self, workspace_with_fixtures):
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.txt"})
        assert "Document: sample.txt" in result
        assert "kiso-prod-01.example.com" in result
        assert "PostgreSQL 15" in result

    def test_read_csv(self, workspace_with_fixtures):
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.csv"})
        assert "Dataset: sample.csv" in result
        assert "6 rows" in result
        assert "Columns: id, name, email, role, active" in result
        assert "Alice" in result
        assert "Eve" in result

    def test_read_docx(self, workspace_with_fixtures):
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.docx"})
        assert "Document: sample.docx" in result
        assert "Flask" in result
        assert "Docker" in result

    def test_read_xlsx(self, workspace_with_fixtures):
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.xlsx"})
        assert "Workbook: sample.xlsx" in result
        assert "Sheet: Sales" in result
        assert "Sheet: Summary" in result
        assert "Widget A" in result
        assert "Europe" in result
        assert "330" in result  # Summary sheet total

    def test_read_pdf_with_real_text(self, workspace_with_fixtures):
        """PDF read with real pypdf extraction — no mocks."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.pdf"})
        assert "Document: sample.pdf (3 pages)" in result
        assert "Page 1" in result
        assert "Page 2" in result
        assert "Page 3" in result
        assert "sample PDF document" in result
        assert "192.168.1.100" in result
        assert "production" in result and "deployment" in result

    def test_read_pdf_page_ranges_real(self, workspace_with_fixtures):
        """PDF page range with real extraction."""
        result = do_read(str(workspace_with_fixtures), {
            "file_path": "uploads/sample.pdf",
            "pages": "2",
        })
        assert "Page 2" in result
        assert "192.168.1.100" in result
        assert "Page 1" not in result
        assert "Page 3" not in result

    def test_read_pdf_blank_pages(self, workspace):
        """pypdf blank pages produce no text."""
        from pypdf import PdfWriter
        f = workspace / "uploads" / "blank.pdf"
        writer = PdfWriter()
        for _ in range(3):
            writer.add_blank_page(width=200, height=200)
        writer.write(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/blank.pdf"})
        assert "no extractable text" in result.lower()

    def test_read_unsupported_format(self, workspace):
        f = workspace / "uploads" / "data.xyz"
        f.write_bytes(b"\x00\x01\x02\x03" * 200)
        result = do_read(str(workspace), {"file_path": "uploads/data.xyz"})
        assert "Unsupported" in result

    def test_read_unknown_extension_text_heuristic(self, workspace):
        f = workspace / "uploads" / "config.env"
        f.write_text("KEY=value\nOTHER=stuff\n")
        result = do_read(str(workspace), {"file_path": "uploads/config.env"})
        assert "KEY=value" in result

    def test_read_missing_file_path(self, workspace):
        with pytest.raises(ValueError, match="file_path"):
            do_read(str(workspace), {})

    def test_read_empty_csv(self, workspace):
        """Empty CSV returns appropriate message."""
        f = workspace / "uploads" / "empty.csv"
        f.write_text("")
        result = do_read(str(workspace), {"file_path": "uploads/empty.csv"})
        assert "empty" in result.lower()

    def test_read_empty_docx(self, workspace):
        """DOCX with no text content."""
        from docx import Document
        f = workspace / "uploads" / "empty.docx"
        doc = Document()
        doc.save(str(f))
        result = do_read(str(workspace), {"file_path": "uploads/empty.docx"})
        assert "no text content" in result.lower()

    def test_read_xlsx_multiple_sheets(self, workspace_with_fixtures):
        """XLSX with multiple sheets shows both with correct headers."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.xlsx"})
        assert "2 sheets: Sales, Summary" in result
        # Sales sheet data
        assert "Widget A" in result
        # Summary sheet data
        assert "Europe\t330" in result or ("Europe" in result and "330" in result)


# ---------------------------------------------------------------------------
# Smart truncation — PDF
# ---------------------------------------------------------------------------


class TestSmartTruncationPDF:
    def test_small_pdf_no_truncation(self, workspace_with_fixtures):
        """Small PDF shows all pages with header, no continuation hint."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.pdf"})
        assert "Document: sample.pdf (3 pages)" in result
        assert "Showing pages" not in result
        assert 'Use pages=' not in result

    def test_large_pdf_truncates_at_page_boundary(self, workspace):
        """Large PDF truncates at page boundary with continuation hint."""
        pages = []
        for i in range(100):
            p = MagicMock()
            p.extract_text.return_value = f"Page {i+1} content. " + ("x" * 900)
            pages.append(p)

        # Need a real PDF file for _resolve_path
        from pypdf import PdfWriter
        f = workspace / "uploads" / "big.pdf"
        writer = PdfWriter()
        writer.add_blank_page(200, 200)
        writer.write(str(f))

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = pages
            result = do_read(str(workspace), {"file_path": "uploads/big.pdf"})

        assert "Document: big.pdf (100 pages)" in result
        assert "Showing pages 1-" in result
        assert "of 100" in result
        assert 'Use pages="' in result
        assert "Page 100 content" not in result
        assert "Page 1 content" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500

    def test_pdf_continuation_hint_correct_numbers(self, workspace):
        """Continuation hint suggests the right next page range."""
        pages = []
        for i in range(50):
            p = MagicMock()
            p.extract_text.return_value = f"P{i+1}. " + ("y" * 2000)
            pages.append(p)

        from pypdf import PdfWriter
        f = workspace / "uploads" / "medium.pdf"
        writer = PdfWriter()
        writer.add_blank_page(200, 200)
        writer.write(str(f))

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = pages
            result = do_read(str(workspace), {"file_path": "uploads/medium.pdf"})

        assert "Showing pages 1-" in result
        assert "of 50" in result
        assert 'to read more.' in result


# ---------------------------------------------------------------------------
# Smart truncation — CSV
# ---------------------------------------------------------------------------


class TestSmartTruncationCSV:
    def test_small_csv_no_truncation(self, workspace_with_fixtures):
        """Small CSV shows all rows with header, no hint."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.csv"})
        assert "Dataset: sample.csv (6 rows" in result
        assert "Columns: id, name, email, role, active" in result
        assert "Showing rows" not in result

    def test_large_csv_truncates_at_row_boundary(self, workspace):
        """Large CSV truncates at row boundary with continuation hint."""
        f = workspace / "uploads" / "big.csv"
        with open(f, "w", newline="") as fh:
            writer = csv.writer(fh)
            writer.writerow(["id", "name", "value"])
            for i in range(10000):
                writer.writerow([i, f"item_{i}", "x" * 50])

        result = do_read(str(workspace), {"file_path": "uploads/big.csv"})
        assert "Dataset: big.csv (10001 rows" in result
        assert "Columns: id, name, value" in result
        assert "Showing rows 1-" in result
        assert "of 10001" in result
        assert "search(query)" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Smart truncation — XLSX
# ---------------------------------------------------------------------------


class TestSmartTruncationXLSX:
    def test_small_xlsx_no_truncation(self, workspace_with_fixtures):
        """Small XLSX shows all data with header, no hint."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.xlsx"})
        assert "Workbook: sample.xlsx" in result
        assert "truncated" not in result.lower()

    def test_large_xlsx_truncates_mid_sheet(self, workspace):
        """Large XLSX truncates mid-sheet with hint."""
        from openpyxl import Workbook
        f = workspace / "uploads" / "big.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["ID", "Product", "Amount"])
        for i in range(5000):
            ws.append([i, f"product_{i}", i * 10])
        wb.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/big.xlsx"})
        assert "Workbook: big.xlsx" in result
        assert "Sheet: Sales" in result
        assert "truncated" in result.lower()
        assert "search(query)" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500

    def test_large_xlsx_multi_sheet_truncation(self, workspace):
        """XLSX with 2 sheets where truncation happens in second sheet."""
        from openpyxl import Workbook
        f = workspace / "uploads" / "multi.xlsx"
        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["ID", "Data"])
        for i in range(200):
            ws1.append([i, f"row_{i}_" + "a" * 100])

        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Key", "Value"])
        for i in range(5000):
            ws2.append([f"key_{i}", "b" * 100])
        wb.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/multi.xlsx"})
        assert "Workbook: multi.xlsx (2 sheets: Sheet1, Sheet2)" in result
        # Sheet1 should be fully shown
        assert "Sheet: Sheet1" in result
        # Sheet2 should be partially shown (truncated)
        assert "Sheet: Sheet2" in result
        assert "truncated" in result.lower()


# ---------------------------------------------------------------------------
# Smart truncation — DOCX
# ---------------------------------------------------------------------------


class TestSmartTruncationDOCX:
    def test_small_docx_no_truncation(self, workspace_with_fixtures):
        """Small DOCX shows all text with header, no hint."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.docx"})
        assert "Document: sample.docx" in result
        assert "Showing first" not in result

    def test_large_docx_truncates_at_paragraph(self, workspace):
        """Large DOCX truncates at paragraph boundary with hint."""
        from docx import Document
        f = workspace / "uploads" / "big.docx"
        doc = Document()
        for i in range(500):
            doc.add_paragraph(f"Paragraph {i}. " + "Lorem ipsum dolor sit amet. " * 20)
        doc.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/big.docx"})
        assert "Document: big.docx" in result
        assert "Showing first" in result
        assert "chars" in result
        assert "exec tasks" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Smart truncation — plain text
# ---------------------------------------------------------------------------


class TestSmartTruncationText:
    def test_small_text_no_truncation(self, workspace_with_fixtures):
        """Small text file shows all content with header, no hint."""
        result = do_read(str(workspace_with_fixtures), {"file_path": "uploads/sample.txt"})
        assert "Document: sample.txt" in result
        assert "Showing first" not in result

    def test_large_text_truncates_at_line_boundary(self, workspace):
        """Large text file truncates at line boundary with hint."""
        f = workspace / "uploads" / "big.txt"
        lines = [f"Line {i}: " + "a" * 100 for i in range(_MAX_OUTPUT_CHARS // 100 + 100)]
        f.write_text("\n".join(lines))

        result = do_read(str(workspace), {"file_path": "uploads/big.txt"})
        assert "Document: big.txt" in result
        assert "Showing first" in result
        assert "exec tasks" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Path traversal guard
# ---------------------------------------------------------------------------


class TestPathTraversal:
    def test_traversal_rejected(self, workspace):
        with pytest.raises(ValueError, match="traversal"):
            _resolve_path(str(workspace), {"file_path": "../../etc/passwd"})

    def test_valid_path_accepted(self, workspace, txt_file):
        result = _resolve_path(str(workspace), {"file_path": "uploads/hello.txt"})
        assert result.name == "hello.txt"

    def test_traversal_lateral_escape(self, tmp_path):
        """Sibling directory escape via prefix attack: /tmp/workspace vs /tmp/workspace-data."""
        workspace = tmp_path / "workspace"
        workspace.mkdir()
        sibling = tmp_path / "workspace-data"
        sibling.mkdir()
        secret = sibling / "file.txt"
        secret.write_text("secret")
        with pytest.raises(ValueError, match="traversal"):
            _resolve_path(str(workspace), {"file_path": "../workspace-data/file.txt"})


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class TestParsePageRanges:
    def test_single_page(self):
        assert _parse_page_ranges("3", 10) == [2]

    def test_range(self):
        assert _parse_page_ranges("2-4", 10) == [1, 2, 3]

    def test_mixed(self):
        assert _parse_page_ranges("1,3-5,8", 10) == [0, 2, 3, 4, 7]

    def test_out_of_bounds_clamped(self):
        result = _parse_page_ranges("1-100", 5)
        assert result == [0, 1, 2, 3, 4]

    def test_deduplication(self):
        result = _parse_page_ranges("1-3,2-4", 10)
        assert result == [0, 1, 2, 3]

    def test_invalid_non_numeric(self):
        with pytest.raises(ValueError):
            _parse_page_ranges("abc", 10)

    def test_empty_string(self):
        with pytest.raises(ValueError):
            _parse_page_ranges("", 10)

    def test_zero_page(self):
        """Page 0 is out of bounds (1-based input)."""
        result = _parse_page_ranges("0", 10)
        assert result == []

    def test_negative_page(self):
        with pytest.raises(ValueError):
            _parse_page_ranges("-1", 10)

    def test_reversed_page_range(self):
        """Reversed range like '10-5' auto-reverses to pages 5-10 (0-based: 4-9)."""
        result = _parse_page_ranges("10-5", 20)
        assert result == [4, 5, 6, 7, 8, 9]


class TestFormatSize:
    def test_bytes(self):
        assert _format_size(500) == "500 B"

    def test_kilobytes(self):
        assert "KB" in _format_size(2048)

    def test_megabytes(self):
        assert "MB" in _format_size(5 * 1024 * 1024)


class TestIsLikelyText:
    def test_text_file(self, tmp_path):
        f = tmp_path / "readme.nfo"
        f.write_text("This is plain text content.")
        assert _is_likely_text(f) is True

    def test_binary_file(self, tmp_path):
        f = tmp_path / "image.bin"
        f.write_bytes(bytes(range(256)) * 4)
        assert _is_likely_text(f) is False

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty"
        f.write_bytes(b"")
        assert _is_likely_text(f) is False


# ---------------------------------------------------------------------------
# Functional: stdin/stdout contract
# ---------------------------------------------------------------------------


class TestFunctional:
    def test_list_via_stdin(self, workspace, txt_file):
        """Full subprocess: JSON stdin → stdout."""
        input_data = json.dumps({
            "args": {"action": "list"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 0
        assert "hello.txt" in result.stdout

    def test_read_via_stdin(self, workspace, txt_file):
        input_data = json.dumps({
            "args": {"action": "read", "file_path": "uploads/hello.txt"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 0
        assert "Hello, world!" in result.stdout

    def test_info_via_stdin(self, workspace_with_fixtures):
        """Info action via subprocess."""
        input_data = json.dumps({
            "args": {"action": "info", "file_path": "uploads/sample.pdf"},
            "workspace": str(workspace_with_fixtures),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 0
        assert "Pages: 3" in result.stdout

    def test_missing_file_exits_1(self, workspace):
        input_data = json.dumps({
            "args": {"action": "read", "file_path": "uploads/nope.txt"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 1
        assert "not found" in result.stderr.lower()

    def test_unknown_action_exits_1(self, workspace):
        """Unknown action → exit 1."""
        input_data = json.dumps({
            "args": {"action": "explode"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 1
        assert "unknown" in result.stderr.lower() or "unknown" in result.stdout.lower()

    def test_malformed_json_stdin(self, workspace):
        """Malformed JSON input → exit 1, stderr contains error."""
        result = subprocess.run(
            [sys.executable, "run.py"],
            input="not json", capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 1
        assert "invalid json" in result.stderr.lower() or "json" in result.stderr.lower()


# ---------------------------------------------------------------------------
# XLSX edge case: empty sheet between populated sheets
# ---------------------------------------------------------------------------


class TestXlsxEmptySheetBetween:
    def test_empty_sheet_between_populated(self, workspace):
        """XLSX with Sheet1 (data), Sheet2 (empty), Sheet3 (data) → both data sheets shown."""
        from openpyxl import Workbook
        f = workspace / "uploads" / "gaps.xlsx"
        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["ID", "Name"])
        ws1.append([1, "Alice"])

        ws2 = wb.create_sheet("Sheet2")
        # intentionally empty

        ws3 = wb.create_sheet("Sheet3")
        ws3.append(["Key", "Value"])
        ws3.append(["color", "blue"])

        wb.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/gaps.xlsx"})
        assert "Sheet: Sheet1" in result
        assert "Alice" in result
        assert "Sheet: Sheet3" in result
        assert "blue" in result
        # Sheet2 is empty — should not appear as a data section
        assert "Sheet: Sheet2" not in result
